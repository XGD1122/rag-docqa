"""Word 文档加载器

使用 python-docx 解析 .docx 文件。
"""

import uuid
from pathlib import Path
from src.loader.base import DocumentLoader
from src.config import Document


class DocxLoader(DocumentLoader):
    """Word 文档加载器"""

    @staticmethod
    def supports(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def load(self, file_path: str) -> Document:
        from docx import Document as DocxDocument

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = DocxDocument(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        return Document(
            doc_id=str(uuid.uuid4()),
            filename=path.name,
            file_type="docx",
            raw_text=full_text,
            metadata={"file_path": str(path.absolute())},
        )
